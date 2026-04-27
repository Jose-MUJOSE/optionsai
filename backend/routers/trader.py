"""
Professional Trader Agent router.

Endpoints:
  POST /api/trader/analyze/{ticker}   — SSE stream of 8 researchers + manager
  POST /api/trader/report             — Generate downloadable Word report
"""
from __future__ import annotations

import io
from typing import Literal, Optional
from datetime import datetime

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel, Field

from backend.services.ai_assistant import client, MODEL
from backend.services.data_fetcher import DataFetcher
from backend.services.ticker_validator import validate_us_ticker
from backend.services.trader_agent import TraderAgentPipeline, RESEARCHER_SPECS

router = APIRouter(tags=["Professional Trader Agent"])

_fetcher = DataFetcher()


def _ensure_us_ticker(ticker: str) -> str:
    t = ticker.upper().strip()
    result = validate_us_ticker(t)
    if not result.valid:
        raise HTTPException(
            status_code=422,
            detail={
                "code": "INVALID_TICKER",
                "message_en": result.reason_en,
                "message_zh": result.reason_zh,
            },
        )
    return t


class TraderAnalyzeRequest(BaseModel):
    mode: Literal["stock", "options"] = Field(
        default="stock",
        description="Whether to analyze the underlying stock or its options",
    )
    locale: Literal["zh", "en"] = Field(default="en")


@router.post("/trader/analyze/{ticker}")
async def trader_analyze(ticker: str, req: TraderAnalyzeRequest = None):
    """
    Run the multi-agent trader pipeline. Streams Server-Sent Events.

    Each event is a JSON object — see TraderAgentPipeline for event schema.
    """
    if req is None:
        req = TraderAnalyzeRequest()

    ticker = _ensure_us_ticker(ticker)

    pipeline = TraderAgentPipeline(client, MODEL, _fetcher)

    async def event_stream():
        async for chunk in pipeline.run(ticker, req.mode, req.locale):
            yield chunk

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ============================================================
# Word report generation
# ============================================================

class TraderReportRequest(BaseModel):
    ticker: str
    mode: Literal["stock", "options"]
    locale: Literal["zh", "en"] = "en"
    researchers: list[dict] = Field(default_factory=list)
    manager: dict = Field(default_factory=dict)


def _build_word_report(req: TraderReportRequest) -> bytes:
    """Generate a Word .docx report and return its bytes."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    is_zh = req.locale == "zh"
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    def heading(text: str, level: int = 1):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.bold = True
        run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def para(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def label_value(label: str, value: str):
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.color.rgb = RGBColor(0, 0, 0)
        run_value = p.add_run(value or "—")
        run_value.font.color.rgb = RGBColor(0, 0, 0)

    # === Title ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        f"OptionsAI 专业交易员分析报告" if is_zh
        else "OptionsAI Professional Trader Analysis Report"
    )
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        f"{req.ticker} · "
        + (("期权分析" if req.mode == "options" else "股票分析") if is_zh
           else ("Options Analysis" if req.mode == "options" else "Stock Analysis"))
        + f" · {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()  # spacer

    # === Final Decision ===
    heading("最终决策" if is_zh else "Final Decision", level=1)
    m = req.manager or {}
    label_value(
        "决策" if is_zh else "Decision",
        str(m.get("decision", "—")),
    )
    label_value(
        "信心度" if is_zh else "Conviction",
        f"{m.get('conviction', '—')}/10",
    )
    if req.mode == "stock":
        label_value("时间周期" if is_zh else "Time Horizon", str(m.get("time_horizon", "—")))
        label_value("入场区间" if is_zh else "Entry Zone", str(m.get("entry_zone", "—")))
        label_value("目标价" if is_zh else "Target Price", str(m.get("target_price", "—")))
        label_value("止损价" if is_zh else "Stop Loss", str(m.get("stop_loss", "—")))
        label_value("仓位建议" if is_zh else "Position Sizing", str(m.get("position_sizing", "—")))
    else:
        label_value("方向" if is_zh else "Direction", str(m.get("direction", "—")))
        label_value("结构" if is_zh else "Structure", str(m.get("structure", "—")))
        label_value("到期日" if is_zh else "Expiration", str(m.get("expiration", "—")))
        label_value("最大亏损" if is_zh else "Max Loss", str(m.get("max_loss", "—")))
        label_value("最大盈利" if is_zh else "Max Profit", str(m.get("max_profit", "—")))
        label_value("盈亏平衡" if is_zh else "Breakeven", str(m.get("breakeven", "—")))
        label_value("胜率" if is_zh else "Win Probability", str(m.get("win_probability", "—")))

    doc.add_paragraph()
    heading("投资逻辑" if is_zh else "Investment Thesis", level=2)
    para(str(m.get("thesis", "")))

    catalysts = m.get("key_catalysts", []) or []
    if catalysts:
        heading("关键催化剂" if is_zh else "Key Catalysts", level=2)
        for c in catalysts:
            doc.add_paragraph(str(c), style="List Bullet")

    risks = m.get("main_risks", []) or []
    if risks:
        heading("主要风险" if is_zh else "Main Risks", level=2)
        for r in risks:
            doc.add_paragraph(str(r), style="List Bullet")

    debate = m.get("debate_summary")
    if debate:
        heading("辩论总结" if is_zh else "Debate Summary", level=2)
        para(str(debate))

    # === Researcher Briefings ===
    doc.add_page_break()
    heading("研究员简报" if is_zh else "Researcher Briefings", level=1)

    for r in req.researchers or []:
        name = r.get("name_zh") if is_zh else r.get("name_en")
        stance = r.get("stance", "neutral")
        confidence = r.get("confidence", "—")
        stance_label = (
            {"bullish": "看多", "bearish": "看空", "neutral": "中性"}[stance]
            if is_zh and stance in ("bullish", "bearish", "neutral")
            else stance
        )
        heading(f"{r.get('icon', '•')} {name}", level=2)
        label_value("立场" if is_zh else "Stance", str(stance_label))
        label_value("信心度" if is_zh else "Confidence", f"{confidence}/10")
        if r.get("headline"):
            label_value("核心观点" if is_zh else "Headline", str(r["headline"]))
        if r.get("evidence"):
            label_value("依据" if is_zh else "Evidence", str(r["evidence"]))

        key_points = r.get("key_points", []) or []
        if key_points:
            sub = doc.add_paragraph()
            run = sub.add_run("关键要点：" if is_zh else "Key Points:")
            run.bold = True
            for kp in key_points:
                doc.add_paragraph(str(kp), style="List Bullet")

        if r.get("risks"):
            label_value("风险" if is_zh else "Risks", str(r["risks"]))

        doc.add_paragraph()

    # Footer disclaimer
    doc.add_page_break()
    disclaimer = (
        "免责声明：本报告由 OptionsAI 自动生成，仅供学习与研究使用。"
        "所有内容均基于公开市场数据，不构成投资建议。投资有风险，决策需谨慎。"
        if is_zh else
        "Disclaimer: This report is auto-generated by OptionsAI for educational and research purposes only. "
        "All content is based on public market data and does not constitute investment advice. "
        "All investments carry risk; please make decisions carefully."
    )
    p = doc.add_paragraph()
    run = p.add_run(disclaimer)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)

    # Serialize to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/trader/report")
async def trader_report(req: TraderReportRequest):
    """Generate a Word .docx report from a completed trader analysis."""
    if not req.ticker:
        raise HTTPException(status_code=400, detail="ticker required")

    try:
        data = _build_word_report(req)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Report generation failed: {type(e).__name__}: {e}")

    filename = f"OptionsAI_Trader_Report_{req.ticker}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
        },
    )


@router.get("/trader/researchers")
async def list_researchers():
    """Return the catalog of researchers (for UI display)."""
    return {
        "researchers": [
            {
                "id": key,
                "name_en": spec["name_en"],
                "name_zh": spec["name_zh"],
                "icon": spec["icon"],
                "color": spec["color"],
            }
            for key, spec in RESEARCHER_SPECS.items()
        ]
    }
