"""
One-shot script: build the Gen AI group project proposal as a .docx.
Plain black text, minimal formatting, no colors.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = r"C:\Users\yuyiz\OneDrive\Desktop\OptionsAI_Proposal.docx"

doc = Document()

# Base style: 11pt Times New Roman, black, single-ish spacing
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0, 0, 0)

# Tight page margins
for section in doc.sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


def heading(text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)


def para(text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def bullet(text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def mono(text: str) -> None:
    """Code / XML style block — monospaced, black."""
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)


# ============================================================
# Title block
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Agentic AI for Retail Options Strategy Recommendation")
tr.bold = True
tr.font.size = Pt(16)
tr.font.color.rgb = RGBColor(0, 0, 0)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Group Project Proposal  |  Generative AI  |  Phase 1")
sr.font.size = Pt(11)
sr.font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph()  # spacer

# ============================================================
# 1. Executive Summary
# ============================================================
heading("1. Executive Summary")
para(
    "We propose an agentic AI system that helps beginner retail investors pick reasonable "
    "options strategies for a single stock. The user gives the system two things: a ticker symbol "
    "and a short view on where they think the stock is going. The agent then decides, step by step, "
    "which pieces of market data to pull, what signals to check, and which one or two options "
    "strategies to recommend. Every reasoning step is logged so the user can see why the agent "
    "reached its conclusion."
)
para(
    "The system is built on the Claude Agent SDK using a ReAct loop with strict XML tagging. "
    "It has six Python tools that return real market and options data from Yahoo Finance, Polygon, "
    "and FINRA. It never places trades, never holds money, and never connects to a brokerage. "
    "Our deliverable is a Colab notebook, a few-shot library of ten gold-standard traces, and an "
    "evaluation suite of twenty scenarios scored on a 1 to 5 rubric."
)

# ============================================================
# 2. Business Problem
# ============================================================
heading("2. Business Problem")
para(
    "Retail investors who want to trade options usually run into two walls. The first is "
    "information overload. A single options decision depends on spot price, implied volatility, "
    "IV rank, the Greeks, the earnings schedule, short interest, analyst targets, and sometimes "
    "unusual options flow. Most beginners do not know which of these matters right now and which "
    "can be ignored. The second wall is strategy complexity. Even after a user understands the "
    "numbers, they still have to choose between long calls, credit spreads, iron condors, and many "
    "other structures, each with its own payoff, break-even, and risk."
)
para(
    "Most free tools only show the raw data and leave the user to figure out the rest. Paid tools "
    "often hand over a recommendation without explaining how they got there. Neither is useful for "
    "a beginner who wants to learn the reasoning, not just copy an answer."
)
para(
    "Our target user is a retail investor who knows what a call and a put are, but does not yet "
    "know when a spread is better than a single option, or how implied volatility changes the "
    "picture before an earnings report. For this user we want the agent to behave like a patient "
    "analyst that looks up the right data, explains its thinking in plain language, and recommends "
    "one or two strategies that fit the situation."
)
para(
    "We keep the scope narrow on purpose. The system only produces analysis and strategy "
    "suggestions. It does not connect to a broker, does not place or simulate real-money trades, "
    "and does not give personalized financial advice."
)

# ============================================================
# 3. Proposed Solution
# ============================================================
heading("3. Proposed Solution")
para(
    "The user types two things into the notebook: a ticker such as AAPL, and a short natural-language "
    "view such as \"I think it goes up slowly over the next month.\" From these two inputs the "
    "agent produces a short structured answer with four parts:"
)
bullet("A summary of the current market context: price, IV level, earnings timing, notable signals.")
bullet("One or two recommended options strategies, with suggested strikes and expiration.")
bullet("For each strategy: the maximum loss, the approximate break-even, and the main risk.")
bullet("A short explanation of why these choices match the user's view.")
para(
    "The key design choice is that the agent does not answer in one big prompt. It reasons in small "
    "steps, and each step is visible. This is important for the grading rubric, and more importantly "
    "it is what makes the answer trustworthy for a beginner who needs to learn."
)

# ============================================================
# 4. System Architecture
# ============================================================
heading("4. System Architecture")
para(
    "The system has four layers: the user interface, the agent controller, the tool layer, and the "
    "data layer."
)
para(
    "The user interface is the Colab notebook. The user enters a ticker and a view, and the "
    "notebook prints the full reasoning trace plus the final answer."
)
para(
    "The agent controller is a Python loop built on the Claude Agent SDK. On every turn the "
    "controller sends the current transcript to the model and asks for one structured block "
    "containing a <thought>, optionally a <call> to a tool, and eventually a <final_answer>. "
    "The controller parses the block, runs the tool if there is one, and appends the result as an "
    "<observation>. The loop stops when the model returns a <final_answer>, or after eight turns, "
    "whichever comes first."
)
para(
    "The tool layer is a set of Python functions registered with the SDK. Each tool has a JSON "
    "schema so the model knows what arguments it takes. The tool layer is the only part of the "
    "system allowed to talk to the outside world. The model itself never fetches data directly."
)
para(
    "The data layer is a set of real providers. Yahoo Finance is used for spot price, OHLCV, "
    "implied volatility, and earnings dates. Polygon.io is used as a secondary source for the "
    "options chain when an API key is available. FINRA RegSHO files are used for daily short "
    "volume. All numbers that reach the model are real; the agent is not allowed to invent data."
)

para("A simplified diagram of the flow:")
mono(
    "+---------------------------+\n"
    "|      User (Colab)         |\n"
    "+-------------+-------------+\n"
    "              |\n"
    "              v\n"
    "+---------------------------+\n"
    "|  Claude Agent Controller  |   <-- Claude Agent SDK\n"
    "|  ReAct loop with XML      |\n"
    "+-------------+-------------+\n"
    "              |   tool_use / result\n"
    "              v\n"
    "+---------------------------+\n"
    "|     Python Tool Layer     |\n"
    "|  (6 registered tools)     |\n"
    "+-------------+-------------+\n"
    "              |\n"
    "              v\n"
    "+---------------------------+\n"
    "|  Real Data Providers      |\n"
    "|  Yahoo / Polygon / FINRA  |\n"
    "+---------------------------+"
)

# ============================================================
# 5. Reasoning Framework
# ============================================================
heading("5. Reasoning Framework: ReAct with XML")
para(
    "We use the ReAct pattern because it maps well to how a human analyst works: think about what "
    "is needed, look up one piece of data, react to what you see, then decide the next step. Each "
    "reasoning step is wrapped in XML tags that the controller parses strictly. If a tag is "
    "missing or malformed, the controller inserts an <error> observation and forces the model to "
    "recover."
)
para("A shortened trace looks like this:")
mono(
    "<thought>\n"
    "  The user expects AAPL to rise slowly over one month. Before picking\n"
    "  a strategy I need the current price, the IV environment, and the\n"
    "  earnings date.\n"
    "</thought>\n"
    "<call>\n"
    '  {"name": "get_market_data", "input": {"ticker": "AAPL"}}\n'
    "</call>\n"
    "<observation>\n"
    '  {"spot": 185.30, "iv_rank": 72, "next_earnings_date": "2026-05-02"}\n'
    "</observation>\n"
    "<thought>\n"
    "  IV rank 72 is high and earnings are in two weeks. A naked long call\n"
    "  would suffer from IV crush. A bull call spread caps that exposure.\n"
    "  I will pull the chain for the expiry after earnings.\n"
    "</thought>\n"
    "<call>\n"
    '  {"name": "get_options_chain", '
    '"input": {"ticker": "AAPL", "expiration": "2026-05-16"}}\n'
    "</call>\n"
    "...\n"
    "<final_answer>\n"
    "  Recommended: 185 / 195 Bull Call Spread, May 16 expiry.\n"
    "  Max loss 4.20, break-even 189.20. Rationale: ...\n"
    "</final_answer>"
)
para(
    "This format does three things at once. It forces a logical internal monologue, it keeps every "
    "tool call machine-parseable, and it produces a reasoning trace that can be read by the grader "
    "without any extra tooling."
)

# ============================================================
# 6. Tools
# ============================================================
heading("6. Tools")
para(
    "The agent has six Python tools registered with the SDK. The rubric only requires three; we "
    "include more because options analysis genuinely needs them and the extra variety gives the "
    "agent real choices to reason about."
)
bullet(
    "get_market_data(ticker): returns spot price, daily change percent, IV rank, 30-day "
    "historical volatility, IV/HV ratio, next earnings date, and the list of available "
    "expirations."
)
bullet(
    "get_options_chain(ticker, expiration): returns the full chain for one expiry. Each row has "
    "strike, bid, ask, mid, implied volatility, open interest, volume, and the Greeks."
)
bullet(
    "get_short_data(ticker): returns FINRA daily short volume for the last twenty trading days "
    "and the most recent short-interest percent of float."
)
bullet("get_news(ticker, limit): returns the most recent headlines for the ticker.")
bullet(
    "get_analyst_targets(ticker): returns the mean analyst price target, the high and low "
    "targets, and the consensus rating."
)
bullet(
    "recommend_strategy(view, context): a deterministic helper that converts a structured view "
    "and a market context into a shortlist of candidate strategies with the payoff math already "
    "computed. The model is still responsible for picking among them, so this helper does not "
    "make the final decision on its own."
)
para(
    "Every tool returns a plain Python dictionary. Every tool wraps its real work in a try/except "
    "block and converts any failure into {\"error\": \"...\"}. This means the agent loop never "
    "crashes on a bad call; the model just sees an error observation and decides what to do next."
)

# ============================================================
# 7. Few-Shot Library
# ============================================================
heading("7. Few-Shot Library")
para(
    "We ship a JSON file with at least ten gold-standard examples. Each example has three fields: "
    "the user input, the full ideal reasoning trace with XML tags, and the final answer. The "
    "examples are loaded into the system prompt at the start of every agent run as in-context "
    "learning material."
)
para("The ten planned examples are:")
bullet("High-IV bullish view, leading to a credit spread.")
bullet("Low-IV bullish view, leading to a long call.")
bullet("Neutral view with earnings nearby, leading to an iron condor.")
bullet("Bearish view with elevated short interest, leading to a put debit spread.")
bullet("A case where one tool returns an error, showing graceful recovery.")
bullet("A vague user input (\"I think something big will happen\"), where the agent asks a clarifying question.")
bullet("A case where the user asks for a strategy that does not fit the view, and the agent pushes back.")
bullet("A liquid ticker example (AAPL).")
bullet("A second liquid ticker example (TSLA) in a different IV environment.")
bullet("A less-liquid ticker example, used to exercise error handling and data-quality caveats.")

# ============================================================
# 8. State, Memory, Error Handling
# ============================================================
heading("8. State, Memory, and Error Handling")
para(
    "The full transcript, including every <thought>, <call>, <observation>, and <error>, is kept "
    "in a single Python list that is passed back to the SDK on each turn. This gives the agent "
    "long context within one session. There is no cross-session memory in this version; every "
    "new question starts from a clean slate."
)
para("Error handling has three layers:")
bullet(
    "At the tool layer, every function catches exceptions and returns a structured error "
    "dictionary instead of raising."
)
bullet(
    "At the controller layer, a malformed XML block or an unknown tool name is converted into an "
    "<error> observation that the model can see and react to."
)
bullet(
    "At the session layer, we cap the loop at eight reasoning turns. If the agent has not "
    "produced a <final_answer> by then, the controller asks it to summarize what it already "
    "knows instead of continuing."
)

# ============================================================
# 9. Platform Rationale
# ============================================================
heading("9. Platform Rationale: Why Claude Agent SDK")
para(
    "We chose the Claude Agent SDK (Option B) over LangGraph for three practical reasons."
)
para(
    "First, the grading rubric gives real weight to the clarity of the agent's internal monologue. "
    "Claude models follow XML-tagged reasoning prompts very reliably, which makes the "
    "<thought>/<call>/<observation> contract easy to enforce in code."
)
para(
    "Second, our problem is not a fixed graph with a small number of branches. The agent decides "
    "how many tools to call and in what order, and that number varies per user input. A manual "
    "tool-use loop maps directly onto this. A StateGraph would force us to define nodes and edges "
    "that we do not actually need."
)
para(
    "Third, options recommendations have low tolerance for hallucination. The Claude SDK makes it "
    "simple to add negative constraints directly in the system prompt — rules like \"never "
    "recommend naked short calls\", \"never invent a price\", \"never claim a guaranteed return\" "
    "— and the model obeys these consistently."
)

# ============================================================
# 10. Evaluation Plan
# ============================================================
heading("10. Evaluation Plan")
para(
    "We build a suite of twenty test scenarios. Ten come from the few-shot library and are treated "
    "as regression tests. Ten more are new stress cases that the model has not seen, including "
    "missing-data cases, contradictory user views, and tickers near earnings."
)
para("Every agent response is scored on five dimensions, each on a 1 to 5 scale:")
bullet(
    "Factual accuracy: does the final answer use only numbers that actually appeared in the tool "
    "observations?"
)
bullet(
    "Strategy fit: does the recommended strategy match the user's view and the IV environment?"
)
bullet("Reasoning clarity: are the <thought> blocks logical, focused, and not redundant?")
bullet("Tool selection: did the agent call the right tools in a reasonable order?")
bullet(
    "Instruction adherence: did the agent respect negative constraints and the analysis-only scope?"
)
para(
    "Scoring is done by at least two human graders and averaged. We also run a small automated "
    "check that flags any final answer containing a number that does not appear in the transcript "
    "of observations, so that simple fabrications are caught even before human review."
)

# ============================================================
# 11. Scope and Disclaimer
# ============================================================
heading("11. Scope and Disclaimer")
para(
    "The system is strictly an analysis and education tool. It recommends options strategies and "
    "explains the reasoning behind them. It does not:"
)
bullet("connect to a brokerage account,")
bullet("place, simulate, or pre-fill real-money trade orders,")
bullet("provide personalized financial, tax, or legal advice,")
bullet("handle funds in any form.")
para(
    "Every final answer ends with a short disclaimer reminding the user that they are responsible "
    "for their own trading decisions and that past data does not guarantee future results."
)

# ============================================================
# 12. Risks and One-Tweak Contingency
# ============================================================
heading("12. Risks and One-Tweak Contingency")
para(
    "The main technical risk is data quality. Free sources occasionally rate-limit us or return "
    "stale values. We reduce this risk by caching recent results in memory and by writing every "
    "tool to degrade gracefully when a provider fails."
)
para(
    "The main scope risk is over-ambition. If by Week 6 the full ReAct loop is unstable on edge "
    "cases, we will use the one-tweak rule to narrow the supported user views from a free-text "
    "input to three categories (bullish, bearish, neutral) and keep everything else the same. "
    "This trims variance without changing the architecture or the grading-relevant pieces."
)

doc.save(OUT_PATH)
print(f"Wrote: {OUT_PATH}")
